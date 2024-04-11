import streamlit as st
import pandas as pd
import requests
import matplotlib.pyplot as plt
import os
from io import BytesIO
from stqdm import stqdm
from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

data_loaded = False

def make_table(model, type, var, aggregation, period, percentile, scenario, model_code, model_calculation, statistic, region_code, region):

    # Building the URL
    url = f'https://cckpapi.worldbank.org/cckp/v1/{model}_{type}_{var}_{type}_{aggregation}_{period}_{percentile}_{scenario}_{model_code}_{model_calculation}_{statistic}/{region_code}?_format=json'

    # Requesting the data
    response = requests.get(url)
    data = response.json()

    table = pd.DataFrame(data['data']).rename_axis('year').reset_index().rename(columns={region_code: f'{var}_{scenario}_{percentile}'})
    table['year'] = pd.to_datetime(table['year'])
    table = table.set_index('year').resample('Y').mean()

    return table

def make_plot_single(table, var, plot=True):
    # Make the plot
    fig = plt.figure(figsize=(6, 3))
    table.plot(ax=plt.gca(), legend=False, color='black', linewidth=0.5, label='Yearly value')
    table.rolling(5, center=True).mean().plot(ax=plt.gca(), legend=False, color='red', linewidth=2. , label='5yr rolling mean')
    ax = plt.gca()

    # hide axis spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    ax.set_ylabel(var_ref[var_ref['Code'] == var].Unit.values[0])
    ax.set_title(var_ref[var_ref['Code'] == var].Variable.values[0])
    ax.grid(True)
    #put legend outside right
    plt.legend(['Yearly value', '5yr rolling mean'], loc='center left', bbox_to_anchor=(1, 0.8), frameon=False)

    if plot:
        st.pyplot(fig)

    return fig

def make_plot_multi(historical, historical_lower, historical_upper,
                    ssp126, ssp126_lower, ssp126_upper,
                    ssp245, ssp245_lower, ssp245_upper,
                    ssp585, ssp585_lower, ssp585_upper,
                    plot=True):

    fig = plt.figure(figsize=(6, 3))
    historical.rolling(5).mean().plot(ax=plt.gca(), color='blue', linewidth=2.)
    historical_lower.rolling(5).mean().plot(ax=plt.gca(), color='blue', linewidth=0.5, linestyle=':')
    historical_upper.rolling(5).mean().plot(ax=plt.gca(), color='blue', linewidth=0.5, linestyle='--')

    ssp126.rolling(5).mean().plot(ax=plt.gca(),  color='green', linewidth=2.)
    ssp126_lower.rolling(5).mean().plot(ax=plt.gca(), color='green', linewidth=0.5, linestyle=':')
    ssp126_upper.rolling(5).mean().plot(ax=plt.gca(),  color='green', linewidth=0.5, linestyle='--')

    ssp245.rolling(5).mean().plot(ax=plt.gca(),  color='orange', linewidth=2.)
    ssp245_lower.rolling(5).mean().plot(ax=plt.gca(), color='orange', linewidth=0.5, linestyle=':')
    ssp245_upper.rolling(5).mean().plot(ax=plt.gca(), color='orange', linewidth=0.5, linestyle='--')

    ssp585.rolling(5).mean().plot(ax=plt.gca(), color='red', linewidth=2. )
    ssp585_lower.rolling(5).mean().plot(ax=plt.gca(), color='red', linewidth=0.5, linestyle=':')
    ssp585_upper.rolling(5).mean().plot(ax=plt.gca(), color='red', linewidth=0.5, linestyle='--')

    ax = plt.gca()

    ax.plot([],[], color='black', linewidth=0, label=' ')
    ax.plot([],[], color='black', linewidth=2., label='Median')
    ax.plot([],[], color='black', linewidth=0.5, linestyle=':', label='Lower')
    ax.plot([],[], color='black', linewidth=0.5, linestyle='--', label='Upper')
    # hide axis spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    ax.set_ylabel(var_ref[var_ref['Code'] == var].Unit.values[0])
    ax.set_title(var_ref[var_ref['Code'] == var].Variable.values[0])
    ax.grid(True)

    #put legend outside right
    handles, labels = ax.get_legend_handles_labels()
    ax.legend(list( handles[i] for i in [0,3,6,9,12,13,14,15] ), ['Historical','SSP 1-2.6','SSP 2-4.5','SSP 5-8.5',' ','Median','Lower','Upper'], loc='center left', bbox_to_anchor=(1, 0.5), frameon=False)

    if plot:
        st.pyplot(fig)

    return fig

def set_up_doc():

    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.15
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading('Climate and Climate Change', level=1)
    title.style.font.color.rgb = RGBColor(0, 0, 0)
    title.bold = True
    title.style.font.name = "Calibri"
    title.style.font.size = Pt(11)
    title.add_run().add_break(WD_BREAK.LINE)

    met = doc.add_heading('Methodology', level=2)
    met.style.font.color.rgb = RGBColor(0, 0, 0)
    met.bold = True
    met.style.font.name = "Calibri"
    met.style.font.size = Pt(11)
    met.add_run().add_break(WD_BREAK.LINE)

    met_desc = [f'The climatic characterization and the analysis of the possible future evolution of the climate for {region}, {country} was carried out through the analysis of the following data:', 
                'For the historical climatic trends, data from the ERA5  (European ReAnalysis version5) reanalysis system were used, which provides hourly estimates of numerous atmospheric, terrestrial and oceanic climatic variables. The data covers the Earth on a 30 km grid and resolves the atmosphere using 137 levels from the surface to an altitude of 80 km. Information on uncertainties is also provided for variables with low spatial and temporal resolutions. Quality assured monthly updates of ERA5 (1950 to present) are released within 3 months in real time. Preliminary daily dataset updates are available to users within 5 days in real time.',
                'Finally, data relating to climate projections for the period 2014-2100 were obtained from the Coupled Model Intercomparison Project Phase 6 (CMIP6)  a project of the Working Group on Coupled Modeling (WGCM) of the World Climate Reserach Program (WGCM), which coordinates since 1995 the global climate modeling experiments carried out by various working groups (for Italy by the Euro-Mediterranean Center for Climate Change (CMCC)), through the definition of common protocols and drivers for all models. The data is made available on a 100x100km grid and for a series of socio-economic scenarios (Shared Socioeconomic Pathways - SSP) which reflect different possible evolution scenarios of greenhouse gas emissions.',
                'The data used are those referring to the Multi model ensemble for the following scenarios:',
                'SSP1-2.6: optimistic scenario in which global CO2 emissions are drastically reduced reaching net zero after 2050 thanks to an evolution of societies towards environmental and social sustainability and temperatures stabilize around 1.8°C more by the end of the century.',
                'SSP2-4.5: Intermediate scenario in which CO2 emissions hover around current levels before starting to decline mid-century but fail to reach net zero by 2100. Socio-economic factors follow their historical trends without significant changes. Progress towards sustainability is slow, with development and income growing unevenly. In this scenario, temperatures rise by 2.7°C by the end of the century.', 
                'SSP5-8.5: Scenario where current CO2 emission levels roughly double by 2050. The global economy is growing rapidly, but this growth is fueled by fossil fuel exploitation and high-intensive lifestyles energy. By 2100, the global average temperature will be as much as 4.4°C higher.']
    
    a = [doc.add_paragraph(met) for met in met_desc]

    reg = doc.add_heading('Regional climatology', level=2)
    reg.style.font.color.rgb = RGBColor(0, 0, 0)
    reg.bold = True
    reg.style.font.name = "Calibri"
    reg.style.font.size = Pt(11)
    reg.add_run().add_break(WD_BREAK.LINE)

    reg_desc = [f'ADD GENERICAL CLIMATE INFORMATION FOR {region}, {country}. Reliable sources are the CCKP, Wikipedia, ...',
                'You can follow the scheme: climate classification of the region according to Kopper']

    a = [doc.add_paragraph(reg) for reg in reg_desc]


    return doc


# make wide screen
st.set_page_config(layout="wide")

st.title('Auto Report Generator - CCKP')

st.write('This app allows you to generate a report based on the data available on the [Climate Change Knowledge Portal (CCKP)](https://climateknowledgeportal.worldbank.org/)')

geo_ref = pd.read_excel('geonames.xlsx', sheet_name='Regions')
var_ref = pd.read_excel('geonames.xlsx', sheet_name='Variables')

era_var_code = ['tas','tasmax','tasmin','tnn','tr','txx','fd','pr','rx1day','rx5day']
era_var = [var_ref[var_ref['Code'] == var].Variable.values[0] for var in era_var_code]

cmip_var_code = ['tas', 'tasmax', 'tasmin', 'tnn', 'tr', 'txx', 'fd','hd30', 'hd35', 'hd40', 'hd45', 'hdd65', 'id',  'cdd65', 'sd',  'tr23', 'tr26', 'tr29', 'pr', 'rx1day', 'rx5day', 'cdd', 'cwd',    'prpercnt', 'r20mm', 'r50mm']
cmip_var = [var_ref[var_ref['Code'] == var].Variable.values[0] for var in cmip_var_code]


# Defining region
st.subheader('Select region')
col1, col2 = st.columns(2)

with col1:
    country = st.selectbox('Country', geo_ref['Country'].unique())
    all_reg = geo_ref[geo_ref['Country'] == country]['State'].unique()

with col2:
    region = st.selectbox('Region', all_reg)
    region_code = geo_ref[(geo_ref['Country'] == country) & (geo_ref['State'] == region)]['State Code'].values[0]

st.write(f'Additional resources about all the countries [here](https://climateknowledgeportal.worldbank.org/general-resources).')
# Defining variable
st.subheader('Select variables')

col1_m, col2_m = st.columns([2,3])

with col1_m:
    with st.expander('ERA5'):
        variable_era = []

        for var_e in era_var:
            with st.container():
                col1, col2= st.columns(2)
                with col1:
                    if st.checkbox(var_e, True, key=f'{var_e}_ERA'):
                        variable_era.append(var_e)
                with col2:
                    st.write(var_ref[var_ref['Variable'] == var_e].Description.values[0])

        variable_era_code = [var_ref[var_ref['Variable'] == var].Code.values[0] for var in variable_era]
    
with col2_m:
    with st.expander('More about ERA5 dataset'):
        st.write(""" 
                 **WARNING** Text generated by Copilot.


                ERA5 is the fifth generation ECMWF atmospheric reanalysis of the global climate.\n
                It covers the period from January 1940 to the present. This comprehensive dataset provides hourly estimates of a wide range of atmospheric, land, and oceanic climate variables.\n
                Let s dive into the details:
                - Spatial Resolution: ERA5 data covers the entire Earth on a 31 km grid.
                - Vertical Resolution: It resolves the atmosphere using 137 levels, extending from the surface up to a height of 80 km.
                - Variable Coverage: ERA5 includes information about uncertainties for all variables, even at reduced spatial and temporal resolutions.
                - Data Availability: You can access ERA5 data for various levels, including single levels, pressure levels, potential temperature levels, and model levels. Additionally, daily and monthly aggregates of the hourly fields are available.
                - Production and Relevance: Produced by the Copernicus Climate Change Service (C3S) at ECMWF, ERA5 combines vast amounts of historical observations into global estimates using advanced modeling and data assimilation systems. It has replaced the previous ERA-Interim reanalysis, which ceased production in August 2019.
                 
                In summary, ERA5 provides a wealth of climate information, making it a valuable resource for understanding our planet s climate dynamics.

                 
                 """)

col1_m, col2_m = st.columns([2,3])

with col1_m:
    with st.expander('CMIP6'):
        variable_cmip = []

        for var_c in cmip_var:
            with st.container():
                col1, col2 = st.columns(2)
                with col1:
                    if var_c in era_var:
                        if st.checkbox(var_c, True, key=f'{var_c}_CMIP'):
                            variable_cmip.append(var_c)
                    else:
                        if st.checkbox(var_c, False, key=f'{var_c}_CMIP'):
                            variable_cmip.append(var_c)
                with col2:
                    st.write(var_ref[var_ref['Variable'] == var_c].Description.values[0])
        
        variable_cmip_code = [var_ref[var_ref['Variable'] == var].Code.values[0] for var in variable_cmip]

with col2_m:
    with st.expander('More about CMIP6 dataset'):
        st.write(""" 
                **WARNING** Text generated by Copilot.
                 
                CMIP6, or the Coupled Model Intercomparison Project Phase 6, represents the next generation of climate models.\n
                These models play a crucial role in helping scientists understand how our climate has changed in the past and how it may evolve in the future.\n
                Let s delve into the details:
                - Purpose and Scope:
                    - Objective: CMIP6 aims to better understand climate changes resulting from natural variability and radiative forcing in a multi-model context.
                    - Participants: Around 100 distinct climate models are being produced across 49 different modelling groups worldwide.
                - Key Features:
                    - Emission Scenarios: CMIP6 explores a wider range of possible future outcomes by running new and updated emission pathways.
                    - Higher Spatial Resolution: CMIP6 models have a higher spatial resolution than their predecessors, allowing them to more realistically represent complex climate processes.
                    - Climate Sensitivity: Some CMIP6 models exhibit notably higher climate sensitivity compared to CMIP5 models. This contributes to projections of greater warming this century, approximately 0.4°C warmer than similar scenarios in CMIP5.
                    - Data Availability: While results from only around 40 CMIP6 models have been published so far, ongoing research may refine these projections as more models become available.
                - Experiments:
                    - CMIP6 models are undertaking various experiments to explore different aspects of climate change. These include scenarios related to future emissions, ocean circulation, carbon cycle feedbacks, and more.
                 """)

# Defining variable
st.subheader('Create report')

doc = set_up_doc()

if st.button('Get data'):

    with st.status("Getting data..."):

        doc.add_page_break() 
        title = doc.add_heading('Historical trends of the main climatic indicators from ERA5', level=2)
        title.style.font.color.rgb = RGBColor(0, 0, 0)
        title.bold = True
        title.style.font.name = "Calibri"
        title.style.font.size = Pt(11)
        title.add_run().add_break(WD_BREAK.LINE)

        st.write('**Getting ERA5 data**')

        for var in stqdm(variable_era_code):

            tab = make_table('era5-x0.5', 'timeseries', var, 'annual', '1950-2020', 'mean', 'historical', 'era5', 'era5', 'mean', region_code, region)
            fig = make_plot_single(tab, var, False)
            fig.savefig('tmp.png', bbox_inches='tight', dpi=300)

            p = doc.add_paragraph().add_run(var_ref[var_ref['Code'] == var].Variable.values[0])

            p.bold = True
            p.italic = True
            doc.add_picture('tmp.png')
            os.remove('tmp.png')
            doc.add_paragraph().add_run(var_ref[var_ref['Code'] == var].Description.values[0])

            # Creating a table object
            table = doc.add_table(rows=1, cols=2)

            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Year'
            row[1].text = f'ERA5 value [{var_ref[var_ref["Code"] == var].Unit.values[0]}]'


            # Adding data from the list to the table
            for y in range(1950,2021,10):

                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(y)
                row[1].text = f'{tab[tab.index.year==y].values[0][0]:6.2f}'
        

            table.style = 'Colorful List'
            doc.add_page_break() 

        
        title = doc.add_heading('Future projections of key climate indicators from CMIP6', level=2)
        title.style.font.color.rgb = RGBColor(0, 0, 0)
        title.bold = True
        title.style.font.name = "Calibri"
        title.style.font.size = Pt(11)
        title.add_run().add_break(WD_BREAK.LINE)

        st.write('**Getting CMIP6 data**')

        for var in stqdm(variable_cmip_code):

            tab_historical = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '1950-2014', 'median', 'historical', 'ensemble', 'all', 'mean', region_code, region)
            tab_historical_lower = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '1950-2014', 'p10', 'historical', 'ensemble', 'all', 'mean', region_code, region)
            tab_historical_upper = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '1950-2014', 'p90', 'historical', 'ensemble', 'all', 'mean', region_code, region)
            
            tab_ssp126 = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'median', 'ssp126', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp126_lower = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p10', 'ssp126', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp126_upper = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p90', 'ssp126', 'ensemble', 'all', 'mean', region_code, region)

            tab_ssp245 = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'median', 'ssp245', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp245_lower = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p10', 'ssp245', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp245_upper = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p90', 'ssp245', 'ensemble', 'all', 'mean', region_code, region)

            tab_ssp585 = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'median', 'ssp585', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp585_lower = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p10', 'ssp585', 'ensemble', 'all', 'mean', region_code, region)
            tab_ssp585_upper = make_table('cmip6-x0.25', 'timeseries', var, 'annual', '2015-2100', 'p90', 'ssp585', 'ensemble', 'all', 'mean', region_code, region)

            tab_tot = pd.concat([tab_historical, tab_historical_lower, tab_historical_upper, tab_ssp126, tab_ssp126_lower, tab_ssp126_upper, tab_ssp245, tab_ssp245_lower, tab_ssp245_upper, tab_ssp585, tab_ssp585_lower, tab_ssp585_upper], axis=1)
            tab_tot = tab_tot.rolling(5).mean()
            
            fig = make_plot_multi(tab_historical, tab_historical_lower, tab_historical_upper,
                            tab_ssp126, tab_ssp126_lower, tab_ssp126_upper,
                            tab_ssp245, tab_ssp245_lower, tab_ssp245_upper,
                            tab_ssp585, tab_ssp585_lower, tab_ssp585_upper, False)
            
            fig.savefig('tmp.png', bbox_inches='tight', dpi=300)

            p = doc.add_paragraph().add_run(var_ref[var_ref['Code'] == var].Variable.values[0])

            p.bold = True
            p.italic = True
            doc.add_picture('tmp.png')
            os.remove('tmp.png')
            doc.add_paragraph().add_run(var_ref[var_ref['Code'] == var].Description.values[0])

            # Creating a table object
            table = doc.add_table(rows=1, cols=4)

            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Year'
            row[1].text = f'SSP 1-2.6 [{var_ref[var_ref["Code"] == var].Unit.values[0]}]'
            row[2].text = f'SSP 2-4.5 [{var_ref[var_ref["Code"] == var].Unit.values[0]}]'
            row[3].text = f'SSP 5-8.5 [{var_ref[var_ref["Code"] == var].Unit.values[0]}]'

            # Adding data from the list to the table
            for y in range(2020,2101,10):
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = str(y)

                row[1].text = f"{tab_tot[tab_tot.index.year==y][f'{var}_ssp126_median'].values[0]:6.2f}"
                row[2].text = f"{tab_tot[tab_tot.index.year==y][f'{var}_ssp245_median'].values[0]:6.2f}"
                row[3].text = f"{tab_tot[tab_tot.index.year==y][f'{var}_ssp585_median'].values[0]:6.2f}"
        

            table.style = 'Colorful List'
            doc.add_page_break() 
        
    st.success('Data loaded successfully')
    data_loaded = True

if data_loaded == True:

    try:
        bio = BytesIO()
        doc.save(bio)

        st.download_button(
                label="Click here to download",
                data=bio.getvalue(),
                file_name=f"{country}_{region}.docx",
                mime="docx"
            )
    
    except:
        pass
